from pathlib import Path
from io import BytesIO

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(r"D:\is\is\P2\plan")
ASSETS = BASE / "_storyboard_assets"
OUT = BASE / "คู่มือและภาพร่างหน้าจอ_Web_Application_เกมมิฟิเคชัน.docx"
DUAL_SCREEN = Path(
    r"C:\Users\phair\AppData\Local\Temp"
    r"\codex-clipboard-f35d894e-5a19-4ac8-93bb-665d13a3d5d5.png"
)

NAVY = "17213D"
PURPLE = "5B45D8"
PURPLE_DARK = "4431B8"
PURPLE_LIGHT = "EEEAFE"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
CYAN = "28A9D7"
MINT = "DFF7EF"
GOLD = "F7C948"
GOLD_LIGHT = "FFF4C7"
INK = "1E2740"
MUTED = "5D6680"
LINE = "DDE2EF"
PAPER = "FFFFFF"
SOFT = "F6F7FC"
RED = "C94B5B"
GREEN = "168466"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        element = borders.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_fixed(table, widths_inches):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), "Tahoma")
    r_fonts.set(qn("w:cs"), "Tahoma")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(paragraph, before=0, after=6, line=1.25, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def clear_para(paragraph):
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_hyperlink(paragraph, text, url, color=PURPLE, underline=True):
    part = paragraph.part
    rid = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Tahoma")
    r_fonts.set(qn("w:cs"), "Tahoma")
    r_pr.append(r_fonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def tag_alt_text(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def crop_image(src, dest, box=None, max_px=2000):
    with Image.open(src) as im:
        im = im.convert("RGB")
        if box:
            im = im.crop(box)
        if im.width > max_px:
            ratio = max_px / im.width
            im = im.resize((max_px, int(im.height * ratio)), Image.Resampling.LANCZOS)
        im.save(dest, quality=92, optimize=True)


def make_cover_art(dest):
    width, height = 1600, 620
    im = Image.new("RGB", (width, height), "#F4F1FF")
    draw = ImageDraw.Draw(im)
    draw.ellipse((1110, -160, 1740, 470), fill="#5B45D8")
    draw.ellipse((1260, 110, 1570, 420), fill="#FFF0A8")
    draw.rounded_rectangle((90, 90, 1030, 525), radius=45, fill="#FFFFFF", outline="#DDD7FA", width=5)
    draw.rounded_rectangle((135, 135, 275, 275), radius=28, fill="#5B45D8")
    draw.text((180, 158), "ก", fill="white", font=ImageFont.truetype(r"C:\Windows\Fonts\tahomabd.ttf", 84))
    draw.text((320, 150), "คำไทยผจญภัย", fill="#18213C", font=ImageFont.truetype(r"C:\Windows\Fonts\tahomabd.ttf", 72))
    draw.text((320, 250), "Web Application เกมมิฟิเคชัน", fill="#5B45D8", font=ImageFont.truetype(r"C:\Windows\Fonts\tahomabd.ttf", 44))
    draw.text((140, 380), "เรียนรู้มาตราตัวสะกดผ่านภารกิจ เกม และผลป้อนกลับทันที", fill="#586078", font=ImageFont.truetype(r"C:\Windows\Fonts\tahoma.ttf", 34))
    for x, y, c in [(1160, 445, "#F7C948"), (1330, 520, "#28A9D7"), (1490, 410, "#FFFFFF")]:
        draw.regular_polygon((x, y, 24), n_sides=5, rotation=0, fill=c)
    im.save(dest, quality=95)


def add_picture_fit(paragraph, path, max_width=6.25, max_height=5.0, alt_title="", alt_desc=""):
    with Image.open(path) as im:
        width_px, height_px = im.size
    ratio = width_px / height_px
    width = max_width
    height = width / ratio
    if height > max_height:
        height = max_height
        width = height * ratio
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width), height=Inches(height))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if alt_title or alt_desc:
        tag_alt_text(shape, alt_title, alt_desc)
    return shape


def add_label(doc, text, fill=PURPLE_LIGHT, color=PURPLE_DARK):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(2.25)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 4, "color": fill},
        bottom={"val": "single", "sz": 4, "color": fill},
        left={"val": "single", "sz": 4, "color": fill},
        right={"val": "single", "sz": 4, "color": fill},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=0)
    r = p.add_run(text)
    set_run_font(r, 9.5, True, color)
    return table


def add_note_box(doc, title, body, fill=SOFT, accent=PURPLE, icon=""):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [0.42, 5.98])
    left, right = table.rows[0].cells
    set_cell_shading(left, accent)
    set_cell_shading(right, fill)
    for cell in (left, right):
        set_cell_border(
            cell,
            top={"val": "single", "sz": 6, "color": LINE},
            bottom={"val": "single", "sz": 6, "color": LINE},
            left={"val": "single", "sz": 6, "color": LINE},
            right={"val": "single", "sz": 6, "color": LINE},
        )
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=0)
    r = p.add_run(icon or "•")
    set_run_font(r, 16, True, PAPER)
    p = right.paragraphs[0]
    set_para_spacing(p, after=2)
    r = p.add_run(title)
    set_run_font(r, 10.5, True, INK)
    p2 = right.add_paragraph()
    set_para_spacing(p2, after=0)
    r2 = p2.add_run(body)
    set_run_font(r2, 9.5, False, MUTED)
    return table


def add_bullets(doc, items, color=INK, font_size=10.5, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        set_para_spacing(p, after=4, line=1.25)
        r = p.add_run(item)
        set_run_font(r, font_size, False, color)


def add_numbered_steps(doc, items):
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [0.55, 5.85])
    for idx, (title, body) in enumerate(items, 1):
        left, right = table.rows[idx - 1].cells
        set_cell_shading(left, PURPLE if idx % 2 else BLUE)
        set_cell_shading(right, PAPER if idx % 2 else SOFT)
        for cell in (left, right):
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": LINE},
                bottom={"val": "single", "sz": 5, "color": LINE},
                left={"val": "single", "sz": 5, "color": LINE},
                right={"val": "single", "sz": 5, "color": LINE},
            )
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, after=0)
        r = p.add_run(str(idx))
        set_run_font(r, 13, True, PAPER)
        p = right.paragraphs[0]
        set_para_spacing(p, after=1)
        r = p.add_run(title)
        set_run_font(r, 10.5, True, INK)
        p2 = right.add_paragraph()
        set_para_spacing(p2, after=0)
        r2 = p2.add_run(body)
        set_run_font(r2, 9.5, False, MUTED)
    return table


def add_screen_meta(doc, screen_id, user, purpose, next_step):
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [1.08, 1.72, 3.60])
    headers = ["รหัสหน้าจอ", "ผู้ใช้งาน", "หน้าที่ของหน้าจอ"]
    values = [screen_id, user, purpose]
    for idx in range(3):
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
        set_cell_shading(table.rows[1].cells[idx], PAPER)
        for row in table.rows:
            set_cell_border(
                row.cells[idx],
                top={"val": "single", "sz": 6, "color": LINE},
                bottom={"val": "single", "sz": 6, "color": LINE},
                left={"val": "single", "sz": 6, "color": LINE},
                right={"val": "single", "sz": 6, "color": LINE},
            )
        p = table.rows[0].cells[idx].paragraphs[0]
        set_para_spacing(p, after=0)
        r = p.add_run(headers[idx])
        set_run_font(r, 9, True, BLUE_DARK)
        p = table.rows[1].cells[idx].paragraphs[0]
        set_para_spacing(p, after=0)
        r = p.add_run(values[idx])
        set_run_font(r, 9.5, idx == 0, INK)
    p = doc.add_paragraph()
    set_para_spacing(p, before=5, after=0)
    r = p.add_run("การเปลี่ยนหน้าจอ  →  ")
    set_run_font(r, 9.5, True, PURPLE)
    r = p.add_run(next_step)
    set_run_font(r, 9.5, False, MUTED)


def title(doc, text, subtitle=None):
    p = doc.add_paragraph(style="Title")
    set_para_spacing(p, before=0, after=5, keep=True)
    r = p.add_run(text)
    set_run_font(r, 22, True, INK)
    if subtitle:
        p2 = doc.add_paragraph()
        set_para_spacing(p2, after=12)
        r2 = p2.add_run(subtitle)
        set_run_font(r2, 10.5, False, MUTED)


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    set_para_spacing(p, before=18, after=10, keep=True)
    r = p.add_run(text)
    set_run_font(r, 16, True, BLUE)
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    set_para_spacing(p, before=14, after=7, keep=True)
    r = p.add_run(text)
    set_run_font(r, 13, True, BLUE)
    return p


def body(doc, text, size=11, color=INK, bold=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_para_spacing(p, after=after)
    r = p.add_run(text)
    set_run_font(r, size, bold, color)
    return p


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=3, after=8)
    r = p.add_run(text)
    set_run_font(r, 9, False, MUTED)
    r.italic = True
    return p


def page_break(doc):
    doc.add_page_break()


def prepare_assets():
    prepared = ASSETS / "prepared"
    prepared.mkdir(exist_ok=True)
    cover = prepared / "cover.jpg"
    make_cover_art(cover)
    crop_image(ASSETS / "01-home.png", prepared / "home.jpg", (0, 0, 1280, 800))
    crop_image(ASSETS / "02-plans.png", prepared / "plans.jpg", (0, 0, 1280, 1260))
    crop_image(ASSETS / "04-teacher.png", prepared / "teacher.jpg")
    crop_image(ASSETS / "05-student.png", prepared / "student.jpg")
    crop_image(ASSETS / "08-game-plan-3.png", prepared / "game3.jpg")
    crop_image(ASSETS / "14-game-activity-plan-3.png", prepared / "activity3.jpg")
    crop_image(ASSETS / "13-game-plan-8.png", prepared / "game8.jpg")
    if DUAL_SCREEN.exists():
        crop_image(DUAL_SCREEN, prepared / "dual.jpg")
    else:
        crop_image(ASSETS / "03-expert.png", prepared / "dual.jpg", (25, 500, 1250, 1800))
    return prepared


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, BLUE_DARK, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Tahoma")
        style._element.rPr.rFonts.set(qn("w:cs"), "Tahoma")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [4.8, 1.6])
    left, right = table.rows[0].cells
    set_cell_shading(left, PAPER)
    set_cell_shading(right, PAPER)
    p = left.paragraphs[0]
    set_para_spacing(p, after=0)
    r = p.add_run("คำไทยผจญภัย  |  คู่มือและ Storyboard")
    set_run_font(r, 8.5, True, PURPLE_DARK)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(p, after=0)
    r = p.add_run("ภาษาไทย ป.2")
    set_run_font(r, 8.5, True, MUTED)
    for cell in (left, right):
        set_cell_border(cell, bottom={"val": "single", "sz": 8, "color": PURPLE})

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=0)
    r = p.add_run("โรงเรียนเทศบาล ๑ (ถนนนครนอก)   •   หน้า ")
    set_run_font(r, 8, False, MUTED)
    add_page_field(p)

    first_header = section.first_page_header
    p = first_header.paragraphs[0]
    clear_para(p)
    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("เอกสารแสดงตัวอย่างนวัตกรรมที่ใช้ในการทดลอง")
    set_run_font(r, 8.5, False, MUTED)


def build_document():
    prepared = prepare_assets()
    doc = Document()
    configure_document(doc)
    cp = doc.core_properties
    cp.title = "คู่มือและภาพร่างหน้าจอ Web Application เกมมิฟิเคชัน คำไทยผจญภัย"
    cp.subject = "ตัวอย่างนวัตกรรมที่ใช้ในการทดลอง สำหรับนักเรียนชั้นประถมศึกษาปีที่ 2"
    cp.author = "ครูไพรัช อินควรชุม"
    cp.keywords = "เกมมิฟิเคชัน, ภาษาไทย, ป.2, มาตราตัวสะกด, Storyboard, คู่มือ"

    # Cover
    p = doc.add_paragraph()
    set_para_spacing(p, after=10)
    r = p.add_run("ชิ้นงานนวัตกรรมที่ใช้ในการทดลอง")
    set_run_font(r, 11, True, PURPLE)
    p = doc.add_paragraph()
    set_para_spacing(p, after=6)
    r = p.add_run("คู่มือการใช้งานเบื้องต้น\nและภาพร่างหน้าจอ (Storyboard)")
    set_run_font(r, 25, True, INK)
    p = doc.add_paragraph()
    set_para_spacing(p, after=12)
    r = p.add_run("Web Application เกมมิฟิเคชัน “คำไทยผจญภัย”")
    set_run_font(r, 16, True, PURPLE)
    p = doc.add_paragraph()
    add_picture_fit(
        p,
        prepared / "cover.jpg",
        6.5,
        3.0,
        "ภาพปกคำไทยผจญภัย",
        "ภาพสัญลักษณ์ Web Application เกมมิฟิเคชันคำไทยผจญภัย",
    )
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [1.65, 4.75])
    cover_data = [
        ("กลุ่มเป้าหมาย", "นักเรียนชั้นประถมศึกษาปีที่ 2"),
        ("หน่วยการเรียนรู้", "มาตราตัวสะกด • 8 แผน • 8 ชั่วโมง"),
        ("ผู้จัดทำ", "ครูไพรัช อินควรชุม"),
        ("สถานศึกษา", "โรงเรียนเทศบาล ๑ (ถนนนครนอก) • ปีการศึกษา 2569"),
    ]
    for row, (label, value) in zip(table.rows, cover_data):
        set_cell_shading(row.cells[0], PURPLE_LIGHT)
        set_cell_shading(row.cells[1], PAPER)
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": LINE},
                bottom={"val": "single", "sz": 5, "color": LINE},
                left={"val": "single", "sz": 5, "color": LINE},
                right={"val": "single", "sz": 5, "color": LINE},
            )
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, 9.5, True, PURPLE_DARK)
        p = row.cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, 9.5, False, INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=8, after=0)
    add_hyperlink(p, "https://webbase-x.github.io/is/P2/", "https://webbase-x.github.io/is/P2/")

    # Overview
    page_break(doc)
    add_label(doc, "ภาพรวมชิ้นงานนวัตกรรม")
    title(doc, "1. Web Application เกมมิฟิเคชัน “คำไทยผจญภัย”", "สื่อหลักสำหรับการจัดการเรียนรู้ภาษาไทย ป.2 เรื่องมาตราตัวสะกด")
    body(
        doc,
        "นวัตกรรมนี้ออกแบบให้ครูควบคุมลำดับการสอนจากหน้าจอเดียว นักเรียนเข้าร่วมด้วยรหัสห้อง "
        "และทำภารกิจบนอุปกรณ์ของตนเอง โดยระบบให้ผลป้อนกลับทันทีและสรุปผลการแข่งขันแบบเรียลไทม์",
    )
    add_note_box(
        doc,
        "จุดมุ่งหมายของนวัตกรรม",
        "ช่วยให้นักเรียนอ่านออกเสียง จำแนกคำตามมาตรา และนำคำไปใช้ได้อย่างมีส่วนร่วม "
        "ขณะเดียวกันครูติดตามความก้าวหน้าและควบคุมคาบเรียนได้เป็นขั้นตอน",
        fill=GOLD_LIGHT,
        accent=GOLD,
        icon="★",
    )
    h2(doc, "ผู้ใช้งานหลัก")
    users = [
        ("ครูผู้สอน", "สร้างห้อง อนุมัตินักเรียน เลือกแผน ควบคุมกิจกรรม เวลา และประกาศผล"),
        ("นักเรียน", "เข้าห้องด้วยรหัส 6 หลัก รับภารกิจ เล่นเกม และเห็นผลป้อนกลับทันที"),
        ("ผู้เชี่ยวชาญ", "ทดลองจอครูและจอนักเรียนในหน้าเดียว เพื่อตรวจคุณภาพสื่อทั้ง 8 แผน"),
    ]
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [1.45, 4.95])
    for idx, (role, desc) in enumerate(users):
        set_cell_shading(table.rows[idx].cells[0], PURPLE if idx == 0 else BLUE if idx == 1 else GREEN)
        set_cell_shading(table.rows[idx].cells[1], PAPER)
        for cell in table.rows[idx].cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": LINE},
                bottom={"val": "single", "sz": 5, "color": LINE},
                left={"val": "single", "sz": 5, "color": LINE},
                right={"val": "single", "sz": 5, "color": LINE},
            )
        p = table.rows[idx].cells[0].paragraphs[0]
        r = p.add_run(role)
        set_run_font(r, 10, True, PAPER)
        p = table.rows[idx].cells[1].paragraphs[0]
        r = p.add_run(desc)
        set_run_font(r, 10, False, INK)
    h2(doc, "องค์ประกอบเกมมิฟิเคชัน")
    add_bullets(
        doc,
        [
            "ภารกิจเป็นลำดับ พร้อมตัวจับเวลาและแถบความก้าวหน้า",
            "คะแนน ดาว และเหรียญรางวัลสำหรับเสริมแรงเชิงบวก",
            "กระดานคะแนนและประกาศผลการแข่งขันหลังจบเกม",
            "ผลป้อนกลับทันทีเมื่อคำตอบถูกหรือผิด",
            "Teacher Dashboard สำหรับควบคุมคาบเรียนและติดตามผู้เล่น",
        ],
    )

    # Flow
    page_break(doc)
    add_label(doc, "Storyboard การใช้งาน")
    title(doc, "2. ลำดับการทำงานของระบบ", "เส้นทางตั้งแต่ครูสร้างห้องจนถึงการสรุปผล")
    add_numbered_steps(
        doc,
        [
            ("เข้าสู่ระบบครู", "ครูเปิด teacher.html และเข้าสู่จอควบคุม"),
            ("สร้างคาบเรียน", "เลือกระดับชั้น/แผน แล้วระบบสร้างรหัสห้องและ QR"),
            ("นักเรียนเข้าห้อง", "นักเรียนกรอกรหัส 6 หลักหรือสแกน QR"),
            ("ตรวจรายชื่อ", "ครูอนุมัติผู้ขอเข้าใหม่จากหน้าจอปัจจุบัน"),
            ("เริ่มกิจกรรม", "ครูนำเสนอสื่อหรือเปิดเกมตามลำดับขั้นของแผน"),
            ("เล่นและรับผลทันที", "นักเรียนทำภารกิจ ระบบให้คะแนนและผลป้อนกลับ"),
            ("ประกาศผล", "เมื่อทุกคนจบหรือหมดเวลา ระบบแสดงอันดับและรูปผู้เล่น"),
        ],
    )
    h2(doc, "วงจรการเรียนรู้แบบเกมมิฟิเคชัน")
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [1.18, 1.18, 1.18, 1.18, 1.68])
    loop = [
        ("1", "รับภารกิจ"),
        ("2", "ลงมือเล่น"),
        ("3", "รับผลทันที"),
        ("4", "สะสมคะแนน"),
        ("5", "เห็นความก้าวหน้า"),
    ]
    for idx, (num, label) in enumerate(loop):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PURPLE if idx in (0, 3) else BLUE if idx in (1, 4) else GREEN)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 5, "color": PAPER},
            bottom={"val": "single", "sz": 5, "color": PAPER},
            left={"val": "single", "sz": 5, "color": PAPER},
            right={"val": "single", "sz": 5, "color": PAPER},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{num}\n{label}")
        set_run_font(r, 9.5, True, PAPER)
    add_note_box(
        doc,
        "หลักการควบคุม",
        "เวลาเป็นตัวช่วยกำกับการสอนเท่านั้น เมื่อหมดเวลาระบบไม่เปลี่ยนหน้าเอง "
        "ครูเป็นผู้เลือก “ขั้นก่อนหน้า / เริ่มเวลาใหม่ / ถัดไป”",
        fill=SOFT,
        accent=CYAN,
        icon="T",
    )

    # S01 Home
    page_break(doc)
    add_label(doc, "S01 • หน้าหลัก")
    title(doc, "3. จุดเริ่มต้นของระบบ", "เลือกเส้นทางเข้าสู่จอครู จอนักเรียน หรือหน้าผู้เชี่ยวชาญ")
    p = doc.add_paragraph()
    add_picture_fit(
        p,
        prepared / "home.jpg",
        6.35,
        4.8,
        "หน้าหลักคำไทยผจญภัย",
        "ภาพหน้าหลักของ Web Application แสดงชื่อระบบและช่องทางเข้าสู่การใช้งาน",
    )
    caption(doc, "ภาพที่ 1 หน้าหลักของ Web Application “คำไทยผจญภัย”")
    add_screen_meta(
        doc,
        "S01",
        "ทุกกลุ่ม",
        "แนะนำระบบและแยกทางเข้าตามบทบาท",
        "ครูไป S02 • นักเรียนไป S03 • ผู้เชี่ยวชาญไป S04",
    )

    # S02 + S03
    page_break(doc)
    add_label(doc, "S02–S03 • เข้าสู่ระบบ")
    title(doc, "4. จอครูและจอนักเรียน", "สองบทบาททำงานร่วมกันผ่านรหัสห้องเรียนเดียวกัน")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [3.15, 3.15])
    for idx, (image_path, heading, desc) in enumerate(
        [
            (prepared / "teacher.jpg", "S02 • จอครู", "เข้าสู่ระบบ สร้างคาบเรียน และควบคุมกิจกรรม"),
            (prepared / "student.jpg", "S03 • จอนักเรียน", "กรอกรหัสห้อง 6 หลักเพื่อเข้าร่วมคาบเรียน"),
        ]
    ):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, SOFT)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 6, "color": LINE},
            bottom={"val": "single", "sz": 6, "color": LINE},
            left={"val": "single", "sz": 6, "color": LINE},
            right={"val": "single", "sz": 6, "color": LINE},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(heading)
        set_run_font(r, 11, True, PURPLE if idx == 0 else BLUE)
        p = cell.add_paragraph()
        add_picture_fit(
            p,
            image_path,
            2.85,
            2.25,
            heading,
            f"ภาพตัวอย่าง{heading}ของระบบคำไทยผจญภัย",
        )
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(desc)
        set_run_font(r, 8.8, False, MUTED)
    h2(doc, "การเชื่อมต่อระหว่างสองจอ")
    add_bullets(
        doc,
        [
            "ครูสร้างรหัสห้องหรือ QR แล้วนำเสนอผ่านจอโปรเจกเตอร์",
            "นักเรียนเข้าห้องและรอการอนุมัติ ครูเห็นคำขอใหม่ได้ทุกขั้นของคาบ",
            "เมื่อครูเริ่มเกม นักเรียนทุกคนได้รับชุดกิจกรรมเดียวกันในรอบนั้น",
        ],
        font_size=10.2,
    )

    # Expert dual screen
    page_break(doc)
    add_label(doc, "S04 • โหมดผู้เชี่ยวชาญ")
    title(doc, "5. ทดลองจอครูและจอนักเรียนพร้อมกัน", "หน้าจำลองอัตราส่วน iPad ใช้งานได้จริงในหน้าเดียว")
    p = doc.add_paragraph()
    add_picture_fit(
        p,
        prepared / "dual.jpg",
        6.4,
        3.25,
        "โหมดผู้เชี่ยวชาญสองจอ",
        "ภาพหน้าจอผู้เชี่ยวชาญแสดงจอครูและจอนักเรียนเคียงกันในหน้าเดียว",
    )
    caption(doc, "ภาพที่ 2 หน้าผู้เชี่ยวชาญสำหรับตรวจสื่อทั้งสองมุมมองพร้อมกัน")
    add_screen_meta(
        doc,
        "S04",
        "ผู้เชี่ยวชาญ / อาจารย์ที่ปรึกษา",
        "ทดลองขั้นตอนจริงและตรวจความสัมพันธ์ระหว่างจอครูกับจอนักเรียน",
        "เลือกแผน → เข้าห้องจำลอง → เล่นเกม → ตรวจผลและบันทึกข้อสังเกต",
    )
    add_note_box(
        doc,
        "ประโยชน์ต่อการประเมินนวัตกรรม",
        "ผู้ประเมินไม่ต้องใช้อุปกรณ์สองเครื่อง สามารถตรวจการควบคุม การตอบสนอง "
        "ความเหมาะสมของเนื้อหา และความต่อเนื่องของกิจกรรมได้ในคราวเดียว",
        fill=MINT,
        accent=GREEN,
        icon="✓",
    )

    # Plans
    page_break(doc)
    add_label(doc, "S05 • แผนการสอน")
    title(doc, "6. ศูนย์รวมแผนการจัดการเรียนรู้", "รองรับ 8 แผน ครบมาตราตัวสะกดตามขอบเขตการทดลอง")
    p = doc.add_paragraph()
    add_picture_fit(
        p,
        prepared / "plans.jpg",
        6.3,
        4.8,
        "หน้ารวมแผนการสอน",
        "ภาพหน้าจอรวมแผนการจัดการเรียนรู้ทั้ง 8 แผน",
    )
    caption(doc, "ภาพที่ 3 หน้ารวมแผนการสอนและทางเข้าสู่สื่อ/เกมของแต่ละแผน")
    add_screen_meta(
        doc,
        "S05",
        "ครู / ผู้เชี่ยวชาญ",
        "เลือกแผน ดาวน์โหลดเอกสาร และเปิดสื่อหรือเกมประจำแผน",
        "เลือกแผนที่ต้องการ → เปิดลำดับการสอนหรือเข้าสู่เกมนักเรียน",
    )

    # Game start
    page_break(doc)
    add_label(doc, "S06 • หน้าเปิดภารกิจ")
    title(doc, "7. ตัวอย่างหน้าตาเกมจริง", "ตัวอย่างแผนที่ 3 มาตราแม่กม")
    p = doc.add_paragraph()
    add_picture_fit(
        p,
        prepared / "game3.jpg",
        6.4,
        3.7,
        "หน้าเปิดภารกิจแม่กม",
        "ภาพหน้าเปิดเกมช่วยแมงมุมตามหาคำแม่กม มีปุ่มฉายสื่อและเข้าเกมนักเรียน",
    )
    caption(doc, "ภาพที่ 4 หน้าเปิดภารกิจ มีเรื่องราว เป้าหมาย และทางเลือกการใช้งาน")
    add_screen_meta(
        doc,
        "S06",
        "ครู / นักเรียน",
        "สร้างบริบทของภารกิจและแยกโหมดฉายสื่อกับโหมดเล่นเกม",
        "ฉายสื่อสอนหน้าชั้น หรือกดเข้าเกมนักเรียนเพื่อเริ่มภารกิจ",
    )
    h2(doc, "องค์ประกอบที่ปรากฏ")
    add_bullets(
        doc,
        [
            "ชื่อภารกิจและมาตราที่เป็นเป้าหมาย",
            "คำอธิบายลักษณะของคำอย่างกระชับ",
            "ปุ่มเข้าสื่อสอน ปุ่มเข้าเกม และทางไปจอคะแนนรวม",
            "คะแนนสะสมและภาพประกอบที่เหมาะกับวัย",
        ],
    )

    # Game activity
    page_break(doc)
    add_label(doc, "S07 • ระหว่างเล่นเกม")
    title(doc, "8. ภารกิจโต้ตอบและผลป้อนกลับทันที", "ตัวอย่าง “กล่องคำแม่กม”")
    p = doc.add_paragraph()
    add_picture_fit(
        p,
        prepared / "activity3.jpg",
        6.4,
        3.7,
        "กิจกรรมกล่องคำแม่กม",
        "ภาพเกมนักเรียนที่ให้ฟังคำ พิจารณาคำ และเลือกใส่กล่องแม่กมภายในเวลาที่กำหนด",
    )
    caption(doc, "ภาพที่ 5 นักเรียนฟังคำ พิจารณาตัวสะกด และตัดสินใจภายในเวลาที่กำหนด")
    add_screen_meta(
        doc,
        "S07",
        "นักเรียน",
        "ฝึกจำแนกคำผ่านการตัดสินใจ พร้อมเวลา คะแนน และผลตอบกลับ",
        "ตอบคำถาม → รับผลทันที → ไปคำถัดไป → สรุปคะแนนเมื่อจบภารกิจ",
    )
    add_note_box(
        doc,
        "หลักการเสริมแรง",
        "คำตอบถูกได้รับคะแนน ดาว เสียง หรือภาพเคลื่อนไหวเชิงบวก "
        "คำตอบผิดได้รับคำแนะนำและมีโอกาสเรียนรู้จากข้อผิดพลาดโดยไม่หยุดการเล่น",
        fill=GOLD_LIGHT,
        accent=GOLD,
        icon="★",
    )

    # Variety
    page_break(doc)
    add_label(doc, "ตัวอย่างความต่อเนื่อง 8 แผน")
    title(doc, "9. รูปแบบเดียวกัน แต่เปลี่ยนธีมตามมาตรา", "ตัวอย่างแผนที่ 8 มาตราแม่กน")
    p = doc.add_paragraph()
    add_picture_fit(
        p,
        prepared / "game8.jpg",
        6.4,
        3.7,
        "หน้าเปิดภารกิจแม่กน",
        "ภาพเกมออกเดินทางมาตราแม่กนในธีมอวกาศ",
    )
    caption(doc, "ภาพที่ 6 หน้าตาเกมแผนที่ 8 ใช้โครงสร้างคุ้นเคยและธีมใหม่เพื่อรักษาความสนใจ")
    h2(doc, "แนวคิดการออกแบบร่วม")
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [1.65, 4.75])
    data = [
        ("คงที่ทุกแผน", "ตำแหน่งสำคัญ การนำทาง คะแนน ความก้าวหน้า และผลป้อนกลับ"),
        ("เปลี่ยนตามเนื้อหา", "คำเป้าหมาย ตัวสะกด เรื่องราว ภาพประกอบ และประเภทภารกิจ"),
        ("รองรับอุปกรณ์", "คอมพิวเตอร์ แท็บเล็ต iPad โทรศัพท์ และจอโปรเจกเตอร์"),
        ("ควบคุมโดยครู", "ครูเลือกเริ่ม หยุด ย้อนกลับ เริ่มเวลาใหม่ หรือไปขั้นถัดไป"),
    ]
    for idx, (label, value) in enumerate(data):
        set_cell_shading(table.rows[idx].cells[0], PURPLE_LIGHT)
        set_cell_shading(table.rows[idx].cells[1], PAPER)
        for cell in table.rows[idx].cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": LINE},
                bottom={"val": "single", "sz": 5, "color": LINE},
                left={"val": "single", "sz": 5, "color": LINE},
                right={"val": "single", "sz": 5, "color": LINE},
            )
        p = table.rows[idx].cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, 9.5, True, PURPLE_DARK)
        p = table.rows[idx].cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, 9.5, False, INK)

    # Teacher guide
    page_break(doc)
    add_label(doc, "คู่มือย่อ • ครูผู้สอน")
    title(doc, "10. วิธีใช้งานสำหรับครู", "ขั้นตอนก่อนสอน ระหว่างสอน และหลังจบเกม")
    add_numbered_steps(
        doc,
        [
            ("เปิดจอควบคุม", "เข้าสู่ teacher.html ด้วยบัญชีที่ได้รับสิทธิ์"),
            ("เลือกแผน", "ตรวจชื่อแผน สื่อ เกม และเวลาของแต่ละขั้น"),
            ("สร้างห้อง", "ระบบแสดงรหัสห้องและ QR สำหรับนักเรียน"),
            ("ตรวจนักเรียน", "อนุมัติชื่อผู้ขอเข้าใหม่ก่อนหรือระหว่างคาบ"),
            ("เริ่มการสอน", "นำเสนอสื่อหรือเริ่มเกมตามลำดับ ครูเป็นผู้ควบคุมการเปลี่ยนหน้า"),
            ("ติดตามผล", "ดูจำนวนผู้เล่น ความคืบหน้า คะแนน และผลตอบแบบเรียลไทม์"),
            ("ประกาศผล", "เมื่อทุกคนเล่นจบหรือหมดเวลา เปิดลำดับประกาศผลการแข่งขัน"),
        ],
    )
    h2(doc, "ปุ่มควบคุมหลัก")
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [2.05, 4.35])
    controls = [
        ("← ขั้นก่อนหน้า", "ย้อนกลับไปยังสื่อหรือกิจกรรมก่อนหน้า"),
        ("↻ เริ่มเวลาใหม่", "ตั้งเวลานับถอยหลังของขั้นปัจจุบันใหม่"),
        ("ถัดไป →", "ไปยังสื่อ เกม หรือประกาศผลลำดับถัดไป"),
        ("ซ่อนรายละเอียด", "ขยายพื้นที่จอโปรเจกเตอร์สำหรับนำเสนอหน้าชั้น"),
    ]
    for idx, (label, desc) in enumerate(controls):
        set_cell_shading(table.rows[idx].cells[0], PURPLE if idx in (0, 2) else BLUE)
        set_cell_shading(table.rows[idx].cells[1], SOFT)
        for cell in table.rows[idx].cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": LINE},
                bottom={"val": "single", "sz": 5, "color": LINE},
                left={"val": "single", "sz": 5, "color": LINE},
                right={"val": "single", "sz": 5, "color": LINE},
            )
        p = table.rows[idx].cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, 9.5, True, PAPER)
        p = table.rows[idx].cells[1].paragraphs[0]
        r = p.add_run(desc)
        set_run_font(r, 9.5, False, INK)

    # Student guide
    page_break(doc)
    add_label(doc, "คู่มือย่อ • นักเรียน")
    title(doc, "11. วิธีใช้งานสำหรับนักเรียน", "คำแนะนำสั้น กระชับ เหมาะกับนักเรียนชั้น ป.2")
    add_numbered_steps(
        doc,
        [
            ("เปิดจอนักเรียน", "เข้า student.html หรือสแกน QR จากจอครู"),
            ("ใส่รหัสห้อง", "กรอกรหัสตัวเลข 6 หลักให้ครบ แล้วกด “ไปต่อ”"),
            ("กรอกชื่อ/เลือกโปรไฟล์", "ใช้ชื่อจริงหรือชื่อที่ครูกำหนด และเลือกรูปประจำตัว"),
            ("รอครูอนุมัติ", "เมื่อตรวจรายชื่อแล้ว ครูจะเปิดกิจกรรมให้พร้อมกัน"),
            ("ทำภารกิจ", "ฟังคำ อ่านคำ ลาก วาง เลือก หรือเรียงคำตามคำสั่ง"),
            ("ดูผลของตนเอง", "รับผลป้อนกลับทันที สะสมดาวและคะแนนจนจบเกม"),
        ],
    )
    h2(doc, "ข้อตกลงในการเล่น")
    add_bullets(
        doc,
        [
            "อ่านคำสั่งให้จบก่อนตอบ และเปิดเสียงเมื่อกิจกรรมต้องฟังคำ",
            "ตอบด้วยตนเอง ไม่ปิดหรือรีเฟรชหน้าจอระหว่างเกม",
            "หากสัญญาณขาด ให้แจ้งครูและเข้าห้องเดิมอีกครั้ง",
            "คะแนนมีไว้แสดงความก้าวหน้าและสร้างกำลังใจในการเรียนรู้",
        ],
    )
    add_note_box(
        doc,
        "การเข้าถึงบนอุปกรณ์",
        "หน้าจอปรับตามขนาดอัตโนมัติ รองรับแนวตั้งและแนวนอน "
        "เพื่อใช้ได้ทั้งโทรศัพท์ แท็บเล็ต iPad และคอมพิวเตอร์",
        fill=MINT,
        accent=GREEN,
        icon="✓",
    )

    # Plans 1-4
    page_break(doc)
    add_label(doc, "สารบัญเกม • แผนที่ 1–4")
    title(doc, "12. เกมมิฟิเคชันประกอบแผนการสอน", "ภาพรวมภารกิจที่ใช้จริงในแต่ละแผน")
    plan_rows_1 = [
        (
            "1\nแม่ ก กา",
            "เพลงมาตราแม่ ก กา • วงล้อเสี่ยงทาย • นักสืบเสียงท้ายคำ • จัดบ้านให้คำ • "
            "รถไฟประโยคแม่ ก กา • บอร์ดโหวตประโยคฮิต • ไขกุญแจหีบสมบัติ",
        ),
        (
            "2\nแม่กง",
            "กล่องคำแม่กง • จรวดประโยคพุ่งทะยาน • แบบทดสอบท้ายบทแม่กง",
        ),
        (
            "3\nแม่กม",
            "กล่องคำแม่กม • ภาพนี้คำอะไร • แบบทดสอบท้ายบทแม่กม",
        ),
        (
            "4\nแม่เกย/แม่เกอว",
            "คู่หู ย–ว • เลือกให้ใช่ • ด่านดาวพิชิตคู่หู ย–ว",
        ),
    ]
    table = doc.add_table(rows=1 + len(plan_rows_1), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [1.35, 5.05])
    table.rows[0].cells[0].text = "แผน"
    table.rows[0].cells[1].text = "เกม/ภารกิจหลัก"
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(
            cell,
            top={"val": "single", "sz": 6, "color": LINE},
            bottom={"val": "single", "sz": 6, "color": LINE},
            left={"val": "single", "sz": 6, "color": LINE},
            right={"val": "single", "sz": 6, "color": LINE},
        )
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 10, True, BLUE_DARK)
    for idx, (plan, games) in enumerate(plan_rows_1, 1):
        set_cell_shading(table.rows[idx].cells[0], PURPLE_LIGHT)
        set_cell_shading(table.rows[idx].cells[1], PAPER if idx % 2 else SOFT)
        for cell in table.rows[idx].cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": LINE},
                bottom={"val": "single", "sz": 5, "color": LINE},
                left={"val": "single", "sz": 5, "color": LINE},
                right={"val": "single", "sz": 5, "color": LINE},
            )
        p = table.rows[idx].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(plan)
        set_run_font(r, 10, True, PURPLE_DARK)
        p = table.rows[idx].cells[1].paragraphs[0]
        r = p.add_run(games)
        set_run_font(r, 10, False, INK)
    h2(doc, "โครงสร้างร่วมของทุกแผน")
    add_bullets(
        doc,
        [
            "ภารกิจเริ่มต้น → ฝึกจำแนก/ประยุกต์ → แบบทดสอบหรือด่านสรุป",
            "นักเรียนทุกคนได้รับชุดคำหรือโจทย์เดียวกันในรอบเดียวกัน",
            "ครูเห็นผลรวมและอันดับหลังจบเกม พร้อมข้อมูลผู้เล่น",
        ],
    )

    # Plans 5-8
    page_break(doc)
    add_label(doc, "สารบัญเกม • แผนที่ 5–8")
    title(doc, "13. เกมมิฟิเคชันประกอบแผนการสอน (ต่อ)", "การผจญภัยต่อเนื่องในมาตราที่มีตัวสะกด")
    plan_rows_2 = [
        (
            "5\nแม่กก",
            "เปิดประตูแม่กก • จริงหรือไม่ ใช่หรือเปล่า • ด่านดาวพิชิตแม่กก",
        ),
        (
            "6\nแม่กด",
            "ล่าสมบัติแม่กด • ถอดรหัสศิลาจารึก • ด่านดาวพิชิตแม่กด",
        ),
        (
            "7\nแม่กบ",
            "เก็บเสบียงแม่กบ • ไขปริศนาชาวเกาะ • ด่านดาวพิชิตแม่กบ",
        ),
        (
            "8\nแม่กน",
            "เติมเชื้อเพลิงยานแม่กน • สแกนรหัสลับต่างดาว • ด่านดาวพิชิตแม่กน",
        ),
    ]
    table = doc.add_table(rows=1 + len(plan_rows_2), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [1.35, 5.05])
    table.rows[0].cells[0].text = "แผน"
    table.rows[0].cells[1].text = "เกม/ภารกิจหลัก"
    set_repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(
            cell,
            top={"val": "single", "sz": 6, "color": LINE},
            bottom={"val": "single", "sz": 6, "color": LINE},
            left={"val": "single", "sz": 6, "color": LINE},
            right={"val": "single", "sz": 6, "color": LINE},
        )
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 10, True, BLUE_DARK)
    for idx, (plan, games) in enumerate(plan_rows_2, 1):
        set_cell_shading(table.rows[idx].cells[0], PURPLE_LIGHT)
        set_cell_shading(table.rows[idx].cells[1], PAPER if idx % 2 else SOFT)
        for cell in table.rows[idx].cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": LINE},
                bottom={"val": "single", "sz": 5, "color": LINE},
                left={"val": "single", "sz": 5, "color": LINE},
                right={"val": "single", "sz": 5, "color": LINE},
            )
        p = table.rows[idx].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(plan)
        set_run_font(r, 10, True, PURPLE_DARK)
        p = table.rows[idx].cells[1].paragraphs[0]
        r = p.add_run(games)
        set_run_font(r, 10, False, INK)
    h2(doc, "กลไกที่ใช้เก็บข้อมูลการทดลอง")
    add_bullets(
        doc,
        [
            "บันทึกการเข้าร่วม ภารกิจที่เล่น และผลคะแนนตามรอบ",
            "แสดงผลการแข่งขันหลังจบเกม ไม่รบกวนสมาธิระหว่างทำภารกิจ",
            "Teacher Dashboard ใช้ติดตามความก้าวหน้ารายบุคคลและภาพรวมชั้นเรียน",
            "หน้าผู้เชี่ยวชาญแยกจากหน้าครู เพื่อทดลองสื่อโดยไม่กระทบข้อมูลจริง",
        ],
    )
    add_note_box(
        doc,
        "ข้อสังเกต",
        "รายการเกมสามารถเพิ่มหรือปรับคำเป้าหมายได้ในอนาคต โดยคงโครงสร้างการใช้งานเดิม "
        "จึงรองรับการพัฒนาต่อเนื่องหลังการทดลอง",
        fill=SOFT,
        accent=CYAN,
        icon="+",
    )

    # Requirements + links
    page_break(doc)
    add_label(doc, "การเตรียมพร้อมและการเข้าถึง")
    title(doc, "14. ข้อกำหนดเบื้องต้นและลิงก์ใช้งาน", "รายการตรวจสอบก่อนนำสื่อไปใช้ในการทดลอง")
    h2(doc, "อุปกรณ์และระบบ")
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [1.60, 4.80])
    requirements = [
        ("อินเทอร์เน็ต", "เชื่อมต่อเครือข่ายที่เสถียรสำหรับครูและนักเรียน"),
        ("เบราว์เซอร์", "Chrome, Edge หรือ Safari เวอร์ชันปัจจุบัน"),
        ("จอครู", "คอมพิวเตอร์หรือ iPad และจอโปรเจกเตอร์/โทรทัศน์"),
        ("จอนักเรียน", "โทรศัพท์ แท็บเล็ต iPad หรือคอมพิวเตอร์หนึ่งเครื่องต่อผู้เล่น/กลุ่ม"),
        ("เสียง", "ลำโพงสำหรับกิจกรรมเพลงและการอ่านออกเสียง"),
    ]
    for idx, (label, value) in enumerate(requirements):
        set_cell_shading(table.rows[idx].cells[0], PURPLE_LIGHT)
        set_cell_shading(table.rows[idx].cells[1], PAPER if idx % 2 == 0 else SOFT)
        for cell in table.rows[idx].cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": LINE},
                bottom={"val": "single", "sz": 5, "color": LINE},
                left={"val": "single", "sz": 5, "color": LINE},
                right={"val": "single", "sz": 5, "color": LINE},
            )
        p = table.rows[idx].cells[0].paragraphs[0]
        r = p.add_run(label)
        set_run_font(r, 9.5, True, PURPLE_DARK)
        p = table.rows[idx].cells[1].paragraphs[0]
        r = p.add_run(value)
        set_run_font(r, 9.5, False, INK)
    h2(doc, "ลิงก์ใช้งาน")
    links = [
        ("หน้าหลัก", "https://webbase-x.github.io/is/P2/"),
        ("จอครู", "https://webbase-x.github.io/is/P2/teacher.html"),
        ("จอนักเรียน", "https://webbase-x.github.io/is/P2/student.html"),
        ("หน้าผู้เชี่ยวชาญ", "https://webbase-x.github.io/is/P2/expert.html"),
        ("แผนการสอน", "https://webbase-x.github.io/is/P2/plans.html"),
    ]
    for label, url in links:
        p = doc.add_paragraph()
        set_para_spacing(p, after=4)
        r = p.add_run(f"{label}: ")
        set_run_font(r, 10, True, INK)
        add_hyperlink(p, url, url)
    add_note_box(
        doc,
        "ความเป็นส่วนตัว",
        "เอกสารนี้ไม่แสดงรหัสผ่านจริง ลิงก์สำหรับผู้ประเมินควรส่งเป็นรายบุคคล "
        "และบัญชีประเมินต้องไม่มีสิทธิ์แก้ไขข้อมูลระบบ",
        fill=GOLD_LIGHT,
        accent=GOLD,
        icon="!",
    )

    # Final
    page_break(doc)
    add_label(doc, "สรุปชิ้นงาน")
    title(doc, "15. ความพร้อมสำหรับใช้เป็นนวัตกรรมในการทดลอง", "เอกสารหนึ่งชิ้นที่แสดงทั้งแนวคิด หน้าตา และวิธีใช้งานจริง")
    body(
        doc,
        "คู่มือและ Storyboard ฉบับนี้แสดงรูปธรรมของ Web Application เกมมิฟิเคชัน “คำไทยผจญภัย” "
        "ตั้งแต่หน้าหลัก การเข้าห้องเรียน การควบคุมของครู การเล่นของนักเรียน โหมดผู้เชี่ยวชาญ "
        "ไปจนถึงตัวอย่างเกมจริงและกลไกเกมมิฟิเคชันที่ใช้ในการทดลอง",
        size=11.5,
    )
    h2(doc, "หลักฐานที่ปรากฏในชิ้นงาน")
    add_bullets(
        doc,
        [
            "ภาพหน้าจอจริงของระบบออนไลน์ ไม่ใช่เพียงคำบรรยายแนวคิด",
            "Storyboard ระบุผู้ใช้ หน้าที่ และการเปลี่ยนไปยังหน้าจอถัดไป",
            "คู่มือย่อสำหรับครูและนักเรียน พร้อมขั้นตอนการใช้งาน",
            "รายการเกมทั้ง 8 แผนและองค์ประกอบเกมมิฟิเคชัน",
            "ข้อกำหนดอุปกรณ์และลิงก์สำหรับทดลองใช้งาน",
        ],
    )
    add_note_box(
        doc,
        "สถานะนวัตกรรม",
        "Web Application พร้อมเปิดทดลองผ่านเว็บ และออกแบบให้รองรับการเพิ่มเติมเนื้อหา "
        "คำเป้าหมาย แผนการสอน และกิจกรรมในอนาคต",
        fill=MINT,
        accent=GREEN,
        icon="✓",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=24, after=6)
    r = p.add_run("คำไทยผจญภัย")
    set_run_font(r, 22, True, PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, after=4)
    r = p.add_run("เรียนภาษาไทยอย่างมีส่วนร่วม ผ่านภารกิจที่เห็นความก้าวหน้าได้")
    set_run_font(r, 11, False, MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_hyperlink(p, "เปิด Web Application", "https://webbase-x.github.io/is/P2/")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
